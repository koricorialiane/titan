import re
from html import escape as html_escape
from pathlib import Path

import markdown
import pandas as pd
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def _format_number(value: float, digits: int = 3) -> str:
    if digits == 0:
        return str(int(round(value)))
    return f"{value:.{digits}f}".rstrip("0").rstrip(".")


def _markdown_cell(value: object) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        return format(value, "g")
    return str(value).replace("|", "\\|")


def _dataframe_to_markdown(df: pd.DataFrame) -> str:
    columns = [str(column) for column in df.columns]
    lines = [
        "| " + " | ".join(columns) + " |",
        "| " + " | ".join("---" for _ in columns) + " |",
    ]
    for row in df.itertuples(index=False, name=None):
        lines.append("| " + " | ".join(_markdown_cell(value) for value in row) + " |")
    return "\n".join(lines)


def build_report_markdown(
    system_table: pd.DataFrame,
    modulation_table: pd.DataFrame,
    scenario_goal_table: pd.DataFrame,
    district_results: dict[str, pd.DataFrame],
    festival_results: dict[str, pd.DataFrame],
    assumptions_table: pd.DataFrame,
    validation_table: pd.DataFrame,
) -> str:
    district_coverage = district_results["coverage"]
    district_capacity = district_results["capacity"]
    festival_coverage = festival_results["coverage"]
    festival_capacity = festival_results["capacity"]
    festival_splitting = festival_results["splitting"]

    district_choice = district_capacity.loc[district_capacity["recommended"]].iloc[0]
    festival_choice = festival_splitting.loc[festival_splitting["recommended"]].iloc[0]

    lines = [
        "# Operacion Nexo 5G/6G - Informe tecnico",
        "",
        "## Resumen",
        "",
        "Este informe desarrolla el reto Operacion Nexo 5G/6G para Nueva Pangea comparando dos escenarios de alta demanda: un distrito financiero urbano y un festival masivo temporal.",
        f"El escenario A queda limitado por capacidad y recomienda N={int(district_choice['reuse_factor_n'])}; el escenario B queda claramente limitado por densidad y exige cell splitting hasta la etapa S{int(festival_choice['split_stage'])}.",
        "",
        "## 1. Introduccion",
        "",
        "El problema de una red movil moderna no se resuelve con un unico numero de cobertura. Una ciudad inteligente obliga a distinguir entre alcance radio, carga de trafico e interferencia. El objetivo del reto es demostrar con base fisica y matematica que la solucion de red debe cambiar cuando cambia la casuistica operativa.",
        "",
        "## 2. Estado del arte",
        "",
        "La evolucion hacia LTE y 5G consolida OFDMA, programacion dinamica de recursos y adaptacion MCS. En paralelo, los modelos de propagacion urbanos y suburbanos siguen siendo la base del dimensionamiento inicial. Para el plano de capacidad, Erlang B sigue siendo una herramienta docente valida cuando se desea estimar bloqueo y area maxima por celda.",
        "",
        "### 2.1 Contexto radio moderno",
        "",
        _dataframe_to_markdown(modulation_table),
        "",
        "## 3. Metodologia",
        "",
        "La metodologia reproduce exactamente la secuencia pedida en la guia docente: ruido, sensibilidad, perdida maxima, radio por cobertura, trafico por usuario, densidad de trafico, canales utiles, capacidad Erlang B y radio por capacidad. El criterio final adopta siempre el menor radio entre cobertura y capacidad.",
        "",
        "### 3.1 Parametros base del sistema",
        "",
        _dataframe_to_markdown(system_table),
        "",
        "### 3.2 Objetivos por escenario",
        "",
        _dataframe_to_markdown(scenario_goal_table),
        "",
        "## 4. Escenarios",
        "",
        "### 4.1 Escenario A - Distrito financiero",
        "",
        f"El distrito financiero parte de {float(district_coverage.iloc[0]['traffic_density_erlang_km2']):.1f} Erl/km2, una SNR requerida de 15 dB y una arquitectura de tres sectores por sitio con analisis de N = 3, 4 y 7.",
        "",
        "### 4.2 Escenario B - Festival global",
        "",
        f"El festival parte de {float(festival_coverage.iloc[0]['traffic_density_erlang_km2']):.1f} Erl/km2, una SNR requerida de 5 dB, celdas omnidireccionales y evaluacion explicita de cell splitting.",
        "",
        "## 5. Resultados",
        "",
        "### 5.1 Cobertura",
        "",
        _dataframe_to_markdown(pd.concat([district_coverage, festival_coverage], ignore_index=True)),
        "",
        "### 5.2 Capacidad en el distrito financiero",
        "",
        _dataframe_to_markdown(district_capacity),
        "",
        "### 5.3 Capacidad y splitting en el festival",
        "",
        _dataframe_to_markdown(festival_capacity),
        "",
        _dataframe_to_markdown(festival_splitting),
        "",
        "## 6. Discusion",
        "",
        f"En el distrito financiero, el radio por cobertura alcanza {float(district_coverage.iloc[0]['coverage_radius_km']):.3f} km, pero el radio por capacidad recomendado cae a {float(district_choice['design_radius_km']):.3f} km. Esto demuestra que la celda se dimensiona por carga y no por alcance.",
        f"En el festival, la diferencia es todavia mas extrema: la cobertura teorica es {float(festival_coverage.iloc[0]['coverage_radius_km']):.3f} km mientras que la capacidad util solo sostiene {float(festival_capacity.iloc[0]['capacity_radius_km']):.3f} km. El problema dominante es la densidad de usuarios.",
        "",
        "## 7. Supuestos y trazabilidad",
        "",
        _dataframe_to_markdown(assumptions_table),
        "",
        "## 8. Validacion frente a la rubrica",
        "",
        _dataframe_to_markdown(validation_table),
        "",
        "## 9. Conclusiones",
        "",
        f"La solucion adoptada para el distrito financiero es N={int(district_choice['reuse_factor_n'])} con tres sectores por sitio. La solucion adoptada para el festival es cell splitting hasta S{int(festival_choice['split_stage'])}. Ambas decisiones se justifican porque son las primeras que satisfacen simultaneamente los criterios de cobertura, capacidad e interferencia bajo las hipotesis declaradas.",
        "",
        "### 9.1 Nueva Pangea 2030",
        "",
        "Si la densidad de usuarios creciera, la evolucion natural del proyecto apuntaria a small cells, mas espectro, beamforming, slicing y edge computing. El codigo del proyecto deja la base cuantitativa para ensayar esas extensiones.",
    ]
    return "\n".join(lines)


def build_calculation_annex_markdown(
    district_results: dict[str, pd.DataFrame],
    festival_results: dict[str, pd.DataFrame],
    assumptions_table: pd.DataFrame,
) -> str:
    district_coverage = district_results["coverage"].iloc[0]
    district_capacity = district_results["capacity"].loc[district_results["capacity"]["recommended"]].iloc[0]
    festival_coverage = festival_results["coverage"].iloc[0]
    festival_capacity = festival_results["capacity"].iloc[0]
    festival_split = festival_results["splitting"].loc[festival_results["splitting"]["recommended"]].iloc[0]

    lines = [
        "# Operacion Nexo 5G/6G - Anexo de calculos",
        "",
        "## Escenario A - Distrito financiero",
        "",
        f"1. Ruido termico: `N = {district_coverage['thermal_noise_dbm']:.3f} dBm`.",
        f"2. Sensibilidad: `Sens = {district_coverage['receiver_sensitivity_dbm']:.3f} dBm`.",
        f"3. Perdida maxima: `Lmax = {district_coverage['max_path_loss_db']:.3f} dB`.",
        f"4. Radio por cobertura: `Rcov = {district_coverage['coverage_radius_km']:.3f} km`.",
        f"5. Trafico por usuario: `Auser = {district_coverage['traffic_per_user_erlang']:.3f} Erl`.",
        f"6. Densidad de trafico: `Akm2 = {district_coverage['traffic_density_erlang_km2']:.3f} Erl/km2`.",
        f"7. Recomendacion final: `N = {int(district_capacity['reuse_factor_n'])}` con `Rdesign = {district_capacity['design_radius_km']:.3f} km`.",
        "",
        "## Escenario B - Festival global",
        "",
        f"1. Ruido termico: `N = {festival_coverage['thermal_noise_dbm']:.3f} dBm`.",
        f"2. Sensibilidad: `Sens = {festival_coverage['receiver_sensitivity_dbm']:.3f} dBm`.",
        f"3. Perdida maxima: `Lmax = {festival_coverage['max_path_loss_db']:.3f} dB`.",
        f"4. Radio por cobertura: `Rcov = {festival_coverage['coverage_radius_km']:.3f} km`.",
        f"5. Radio por capacidad: `Rcap = {festival_capacity['capacity_radius_km']:.3f} km`.",
        f"6. Etapa de splitting recomendada: `S{int(festival_split['split_stage'])}` con `R = {festival_split['radius_km']:.3f} km`.",
        "",
        "## Supuestos aplicados",
        "",
        _dataframe_to_markdown(assumptions_table),
    ]
    return "\n".join(lines)


def build_defense_brief_markdown(
    district_results: dict[str, pd.DataFrame],
    festival_results: dict[str, pd.DataFrame],
) -> str:
    district_capacity = district_results["capacity"].loc[district_results["capacity"]["recommended"]].iloc[0]
    festival_split = festival_results["splitting"].loc[festival_results["splitting"]["recommended"]].iloc[0]
    return "\n".join(
        [
            "# Operacion Nexo 5G/6G - Guion de defensa",
            "",
            "- El reto separa de forma explicita cobertura y capacidad.",
            f"- En el distrito financiero la recomendacion es N={int(district_capacity['reuse_factor_n'])} con tres sectores por sitio.",
            f"- En el festival la cobertura no es el problema; la solucion es splitting hasta S{int(festival_split['split_stage'])}.",
            "- El criterio final siempre adopta el radio mas restrictivo.",
            "- Todos los resultados estan trazados en tablas, graficas y supuestos declarados.",
        ]
    )


def build_html_document(document_title: str, markdown_text: str) -> str:
    html_body = markdown.markdown(markdown_text, extensions=["tables"])
    return f"""<!doctype html>
<html lang=\"es\">
<head>
  <meta charset=\"utf-8\">
  <meta name=\"viewport\" content=\"width=device-width, initial-scale=1\">
  <title>{html_escape(document_title)}</title>
  <style>
    body {{ font-family: Arial, Helvetica, sans-serif; margin: 0; background: #f8f5ef; color: #1f242c; }}
    main {{ max-width: 1100px; margin: 32px auto; background: white; padding: 36px 42px; border-radius: 16px; box-shadow: 0 12px 34px rgba(0,0,0,0.08); }}
    h1, h2, h3 {{ color: #1a1f26; }}
    table {{ width: 100%; border-collapse: collapse; margin: 1rem 0 1.4rem 0; font-size: 0.94rem; }}
    th, td {{ border: 1px solid #ddd5ca; padding: 0.55rem 0.6rem; text-align: left; vertical-align: top; }}
    th {{ background: #f5eee4; color: #c85f3f; text-transform: uppercase; font-size: 0.78rem; }}
    code {{ background: #f4ede4; padding: 0.12rem 0.28rem; border-radius: 4px; }}
  </style>
</head>
<body>
  <main>{html_body}</main>
</body>
</html>
"""


def _strip_markdown(text: str) -> str:
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = text.replace("**", "").replace("*", "")
    text = re.sub(r"^#+\s*", "", text)
    return text


def write_html_document(output_path: Path, document_title: str, markdown_text: str) -> None:
    output_path.write_text(build_html_document(document_title, markdown_text), encoding="utf-8")


def write_pdf_document(output_path: Path, document_title: str, markdown_text: str, asset_root: Path) -> None:
    styles = getSampleStyleSheet()
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=9.5, leading=12)
    heading = ParagraphStyle("Heading", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=13, leading=16)
    title_style = ParagraphStyle("Title", parent=styles["Title"], fontName="Helvetica-Bold", fontSize=18, leading=22)

    doc = SimpleDocTemplate(str(output_path), title=document_title, pagesize=A4, leftMargin=1.7 * cm, rightMargin=1.7 * cm, topMargin=1.8 * cm, bottomMargin=1.8 * cm)
    story = [Paragraph(html_escape(document_title), title_style), Spacer(1, 0.35 * cm)]

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.18 * cm))
            continue
        if stripped.startswith("#"):
            story.append(Paragraph(html_escape(_strip_markdown(stripped)), heading))
        else:
            story.append(Paragraph(html_escape(_strip_markdown(stripped)), body))
    doc.build(story)


def write_docx_document(output_path: Path, document_title: str, markdown_text: str, asset_root: Path) -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("# "):
            document.add_heading(_strip_markdown(stripped), level=0)
        elif stripped.startswith("## "):
            document.add_heading(_strip_markdown(stripped), level=1)
        elif stripped.startswith("### "):
            document.add_heading(_strip_markdown(stripped), level=2)
        elif stripped.startswith("- "):
            document.add_paragraph(_strip_markdown(stripped[2:]), style="List Bullet")
        else:
            document.add_paragraph(_strip_markdown(stripped))
    document.save(str(output_path))
